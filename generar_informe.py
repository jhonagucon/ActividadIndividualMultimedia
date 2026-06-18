"""
generar_informe.py
Genera el Informe Técnico del Trabajo Individual de Multimedia.
Autor: Jonathan Gutierrez Condori — UMSA Informática 2026
"""

import os
import sys

try:
    sys.stdout.reconfigure(encoding='utf-8')
except AttributeError:
    pass

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def crear_informe():
    doc = Document()

    # ── Márgenes ──────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Paleta UMSA ───────────────────────────────────────────
    C_NAVY   = RGBColor(31,  73, 125)   # Azul marino UMSA
    C_INDIGO = RGBColor(55,  65, 181)   # Índigo secundario
    C_DARK   = RGBColor(30,  30,  30)   # Cuerpo de texto
    C_MUTED  = RGBColor(100, 100, 100)  # Texto secundario
    C_GREEN  = RGBColor(21, 128,  61)   # Verde acento (código)
    C_BG_COD = 'F0FDF4'                 # Fondo bloque de código
    C_BG_TAB = 'EFF6FF'                 # Fondo tabla

    # ── Helpers ───────────────────────────────────────────────
    def cell_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=160, right=160):
        tcPr  = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
            node = OxmlElement(m)
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def draw_hr(color='1F497D', sz='6'):
        p   = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    sz)
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), color)
        pBdr.append(bot)
        pPr.append(pBdr)

    def titulo_h1(texto, num=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(22)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.keep_with_next = True
        t = f"{num}. {texto}" if num else texto
        run = p.add_run(t)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(15)
        run.bold       = True
        run.font.color.rgb = C_NAVY
        return p

    def titulo_h2(texto):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(texto)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(12)
        run.bold       = True
        run.font.color.rgb = C_INDIGO
        return p

    def titulo_h3(texto):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(texto)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(11)
        run.bold       = True
        run.font.color.rgb = C_NAVY
        return p

    def parrafo(texto, space_after=5, bold=False, italic=False, indent=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if indent:
            p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(texto)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10.5)
        run.font.color.rgb = C_DARK
        run.bold   = bold
        run.italic = italic
        return p

    def bullet(texto, level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent  = Cm(0.8 + level * 0.6)
        run = p.add_run(texto)
        run.font.name  = 'Calibri'
        run.font.size  = Pt(10.5)
        run.font.color.rgb = C_DARK

    def agregar_codigo(codigo_str, lang_label=""):
        # Encabezado del bloque
        tbl_h = doc.add_table(rows=1, cols=1)
        tbl_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ch = tbl_h.cell(0, 0)
        cell_bg(ch, '1F497D')
        set_cell_margins(ch, top=60, bottom=60, left=160, right=160)
        ph = ch.paragraphs[0]
        rh = ph.add_run(lang_label or "Código")
        rh.font.name  = 'Calibri'
        rh.font.size  = Pt(8)
        rh.bold       = True
        rh.font.color.rgb = RGBColor(255, 255, 255)

        # Cuerpo del bloque
        tbl_b = doc.add_table(rows=1, cols=1)
        tbl_b.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cb = tbl_b.cell(0, 0)
        cell_bg(cb, C_BG_COD)
        set_cell_margins(cb, top=140, bottom=140, left=200, right=200)

        tcPr = cb._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        left_b = OxmlElement('w:left')
        left_b.set(qn('w:val'),   'single')
        left_b.set(qn('w:sz'),    '20')
        left_b.set(qn('w:space'), '0')
        left_b.set(qn('w:color'), '15803D')
        tcBorders.append(left_b)
        for side in ['top', 'bottom', 'right']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
        tcPr.append(tcBorders)

        pb = cb.paragraphs[0]
        pb.paragraph_format.space_after  = Pt(0)
        pb.paragraph_format.line_spacing = 1.0
        rb = pb.add_run(codigo_str.strip())
        rb.font.name  = 'Courier New'
        rb.font.size  = Pt(8.5)
        rb.font.color.rgb = C_GREEN

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def agregar_figura(ruta_img, descripcion, ancho_in=5.5):
        if os.path.exists(ruta_img):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(6)
            p_img.paragraph_format.space_after  = Pt(3)
            p_img.add_run().add_picture(ruta_img, width=Inches(ancho_in))
            p_desc = doc.add_paragraph()
            p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_desc.paragraph_format.space_after = Pt(8)
            rdesc = p_desc.add_run(descripcion)
            rdesc.font.name  = 'Calibri'
            rdesc.font.size  = Pt(9)
            rdesc.font.color.rgb = C_MUTED
            rdesc.italic = True
        else:
            parrafo(f"[Figura no disponible: {ruta_img}]", italic=True)

    # ════════════════════════════════════════════════════════════
    # CARÁTULA
    # ════════════════════════════════════════════════════════════
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    for txt, sz, bold, color in [
        ("UNIVERSIDAD MAYOR DE SAN ANDRÉS", 15, True,  C_NAVY),
        ("FACULTAD DE CIENCIAS PURAS Y NATURALES", 11, True,  C_INDIGO),
        ("CARRERA DE INFORMÁTICA", 11, True,  C_INDIGO),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(txt)
        r.font.name = 'Calibri'; r.font.size = Pt(sz); r.bold = bold; r.font.color.rgb = color

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    logo_path = "logo_umsa.png"
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(20)
        p_logo.add_run().add_picture(logo_path, width=Inches(1.9))
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(30)

    draw_hr('1F497D', '12')

    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t.paragraph_format.space_before = Pt(14)
    p_t.paragraph_format.space_after  = Pt(8)
    rt = p_t.add_run("INFORME TÉCNICO DE TRABAJO INDIVIDUAL")
    rt.font.name = 'Calibri'; rt.font.size = Pt(18); rt.bold = True; rt.font.color.rgb = C_NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    rs = p_sub.add_run("Multimedia — Procesamiento de Imágenes y Producción Audiovisual")
    rs.font.name = 'Calibri'; rs.font.size = Pt(12); rs.italic = True; rs.font.color.rgb = C_INDIGO

    draw_hr('1F497D', '12')
    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Tabla de datos del estudiante
    tbl = doc.add_table(rows=5, cols=2)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl.style = 'Table Grid'
    datos = [
        ("Estudiante",   "JONATHAN GUTIERREZ CONDORI"),
        ("Materia",      "Multimedia"),
        ("Carrera",      "Informática"),
        ("Institución",  "Universidad Mayor de San Andrés (UMSA)"),
        ("Gestión",      "2026"),
    ]
    for i, (label, valor) in enumerate(datos):
        row = tbl.rows[i]
        cell_bg(row.cells[0], 'DBEAFE')
        cell_bg(row.cells[1], 'F8FAFC')
        set_cell_margins(row.cells[0])
        set_cell_margins(row.cells[1])
        rl = row.cells[0].paragraphs[0].add_run(label)
        rl.font.name = 'Calibri'; rl.font.size = Pt(10.5); rl.bold = True; rl.font.color.rgb = C_NAVY
        rv = row.cells[1].paragraphs[0].add_run(valor)
        rv.font.name = 'Calibri'; rv.font.size = Pt(10.5); rv.font.color.rgb = C_DARK

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════
    # 1. INTRODUCCIÓN
    # ════════════════════════════════════════════════════════════
    titulo_h1("INTRODUCCIÓN", num=1)
    draw_hr()
    parrafo(
        "El presente informe documenta el desarrollo del Trabajo Individual de la materia de Multimedia "
        "de la Carrera de Informática, UMSA. El proyecto abarca tres actividades individuales de "
        "procesamiento de imágenes digitales y producción multimedia, implementadas en Python con librerías "
        "especializadas y presentadas en una plataforma web interactiva."
    )
    parrafo(
        "Cada actividad trabaja directamente a nivel de píxel o de señal, sin depender de algoritmos de alto "
        "nivel predefinidos, demostrando el dominio de los fundamentos del procesamiento de imágenes y la "
        "producción audiovisual."
    )

    # ════════════════════════════════════════════════════════════
    # 2. OBJETIVOS
    # ════════════════════════════════════════════════════════════
    titulo_h1("OBJETIVOS", num=2)
    draw_hr()
    titulo_h2("2.1 Objetivo General")
    parrafo(
        "Desarrollar e implementar algoritmos de procesamiento de imágenes a nivel de píxel y producir un "
        "cover multimedia, integrando los resultados en una plataforma web funcional."
    )
    titulo_h2("2.2 Objetivos Específicos")
    bullet("Clasificar superficies (césped, tierra, cemento, asfalto) usando análisis de color HSV y varianza local.")
    bullet("Implementar un filtro de promedio con ventana deslizante de 3×3 píxeles para reducir ruido.")
    bullet("Producir el cover de 'La Vaca Lola' sincronizando audio TTS con una animación de video.")
    bullet("Publicar el proyecto completo en un repositorio GitHub con documentación.")
    bullet("Desplegar una plataforma web que demuestre los algoritmos en tiempo real.")

    # ════════════════════════════════════════════════════════════
    # 3. HERRAMIENTAS Y TECNOLOGÍAS
    # ════════════════════════════════════════════════════════════
    titulo_h1("HERRAMIENTAS Y TECNOLOGÍAS UTILIZADAS", num=3)
    draw_hr()
    parrafo(
        "A continuación se detallan todas las herramientas, lenguajes y librerías utilizadas en el desarrollo "
        "de cada actividad del trabajo individual."
    )

    # Tabla de herramientas
    tbl2 = doc.add_table(rows=9, cols=3)
    tbl2.style = 'Table Grid'
    encabezados = ["Herramienta / Librería", "Versión / Fuente", "Uso Principal"]
    herramientas = [
        ("Python 3.11+",         "python.org",          "Lenguaje principal de todos los scripts"),
        ("OpenCV (cv2)",         "pip install opencv-python", "Lectura/escritura de imágenes, operaciones matriciales"),
        ("NumPy",                "pip install numpy",   "Manipulación de arrays de píxeles, cálculo vectorizado"),
        ("edge-tts",             "pip install edge-tts","Síntesis de voz TTS (Microsoft Neural Voices) para el cover"),
        ("FFmpeg",               "ffmpeg.org",          "Composición de video + audio, render del cover final"),
        ("Pillow (PIL)",         "pip install Pillow",  "Procesamiento adicional de imágenes, conversión de formatos"),
        ("Three.js (r128)",      "cdnjs.cloudflare.com","Partículas 3D animadas en la cabecera del sitio web"),
        ("Canvas API (HTML5)",   "Nativa en navegador", "Procesamiento de imágenes en vivo directamente en la web"),
    ]
    row0 = tbl2.rows[0]
    cell_bg(row0.cells[0], '1F497D'); cell_bg(row0.cells[1], '1F497D'); cell_bg(row0.cells[2], '1F497D')
    for ci, enc in enumerate(encabezados):
        r = row0.cells[ci].paragraphs[0].add_run(enc)
        r.font.name = 'Calibri'; r.font.size = Pt(10); r.bold = True; r.font.color.rgb = RGBColor(255,255,255)
        set_cell_margins(row0.cells[ci])
    for i, (h, v, uso) in enumerate(herramientas):
        row = tbl2.rows[i + 1]
        bg = 'F8FAFC' if i % 2 == 0 else 'EFF6FF'
        cell_bg(row.cells[0], bg); cell_bg(row.cells[1], bg); cell_bg(row.cells[2], bg)
        for ci, txt in enumerate([h, v, uso]):
            r = row.cells[ci].paragraphs[0].add_run(txt)
            r.font.name = 'Calibri'; r.font.size = Pt(9.5)
            r.bold = (ci == 0)
            r.font.color.rgb = C_NAVY if ci == 0 else C_DARK
            set_cell_margins(row.cells[ci])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ════════════════════════════════════════════════════════════
    # 4. ACTIVIDAD A — CLASIFICACIÓN DE TEXTURAS
    # ════════════════════════════════════════════════════════════
    doc.add_page_break()
    titulo_h1("ACTIVIDAD A — CLASIFICACIÓN DE TEXTURAS", num=4)
    draw_hr()

    titulo_h2("4.1 Descripción del Problema")
    parrafo(
        "Se requiere desarrollar una aplicación capaz de diferenciar tipos de superficies dentro de una imagen "
        "(césped, tierra, cemento y asfalto) aplicando principios similares a los utilizados en clasificación "
        "de imágenes satelitales. El algoritmo debe operar directamente sobre la matriz de píxeles, sin usar "
        "clasificadores de alto nivel como redes neuronales."
    )

    titulo_h2("4.2 Fundamento Teórico")
    titulo_h3("Espacio de Color HSV")
    parrafo(
        "El espacio HSV (Hue-Saturation-Value) separa la información de color (tono H) de la luminosidad (V) "
        "y la pureza del color (S). Esto es más robusto que RGB para clasificar texturas porque es menos "
        "sensible a cambios de iluminación. La conversión se realiza sobre cada píxel individualmente:"
    )
    agregar_codigo(
        """# Conversión RGB → HSV (sobre cada píxel)
def rgb_a_hsv(r, g, b):
    r, g, b = r/255.0, g/255.0, b/255.0
    cmax = max(r, g, b)        # Valor (V)
    cmin = min(r, g, b)
    delta = cmax - cmin        # Diferencia para calcular H

    # Tono (H) en grados 0°–360°
    if delta == 0:
        h = 0
    elif cmax == r:
        h = 60 * (((g - b) / delta) % 6)
    elif cmax == g:
        h = 60 * (((b - r) / delta) + 2)
    else:
        h = 60 * (((r - g) / delta) + 4)

    s = 0 if cmax == 0 else (delta / cmax) * 100  # Saturación 0–100
    v = cmax * 255                                  # Valor 0–255
    return h, s, v""",
        "Python — Conversión RGB → HSV (píxel a píxel)"
    )

    titulo_h3("Varianza Local de Brillo")
    parrafo(
        "La varianza en una vecindad de 5×5 mide la rugosidad o textura local del píxel. Una superficie "
        "uniforme (como cemento liso) tendrá baja varianza, mientras que una textura rugosa (asfalto con "
        "piedras) tendrá alta varianza:"
    )
    agregar_codigo(
        """# Varianza de brillo en ventana 5×5
def varianza_5x5(img_gray, x, y):
    H, W = img_gray.shape
    vecinos = []
    for ky in range(-2, 3):       # -2, -1, 0, +1, +2
        for kx in range(-2, 3):
            nx, ny = x + kx, y + ky
            if 0 <= nx < W and 0 <= ny < H:
                vecinos.append(float(img_gray[ny, nx]))
    mean = sum(vecinos) / len(vecinos)
    var  = sum((v - mean)**2 for v in vecinos) / len(vecinos)
    return var""",
        "Python — Varianza local 5×5"
    )

    titulo_h2("4.3 Algoritmo de Clasificación Completo")
    parrafo(
        "Con el tono H, la saturación S, el valor V y la varianza calculados para cada píxel, "
        "se aplican las siguientes reglas de clasificación:"
    )
    agregar_codigo(
        """import cv2
import numpy as np

img_bgr = cv2.imread('paisaje.jpg')
img_rgb = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2RGB)
H, W, _ = img_rgb.shape

# Convertir toda la imagen a escala de grises para calcular varianza
img_gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY).astype(float)

# Mapa de salida con colores de clasificación
resultado = np.zeros((H, W, 3), dtype=np.uint8)

# Colores de clasificación
COLORES = {
    'cesped':  (34,  180,  34),   # Verde
    'tierra':  (160,  90,  42),   # Marrón
    'cemento': (200, 200, 200),   # Gris claro
    'asfalto': ( 72,  72,  72),   # Gris oscuro
    'otros':   (128,   0, 128),   # Morado
}

for y in range(H):
    for x in range(W):
        r, g, b = img_rgb[y, x]
        h, s, v = rgb_a_hsv(r, g, b)           # Convertir a HSV
        var = varianza_5x5(img_gray, x, y)      # Varianza 5×5

        # Reglas de clasificación basadas en HSV + varianza
        if 30 <= h <= 90 and s >= 40 and var > 5:
            clase = 'cesped'
        elif 5 <= h < 30 and s >= 30 and var > 15:
            clase = 'tierra'
        elif s < 30 and v >= 100 and var <= 15:
            clase = 'cemento'
        elif s < 35 and v < 100 and var > 15:
            clase = 'asfalto'
        else:
            clase = 'otros'

        resultado[y, x] = COLORES[clase]

cv2.imwrite('resultado_texturas.jpg', cv2.cvtColor(resultado, cv2.COLOR_RGB2BGR))
print("✅ Clasificación completada → resultado_texturas.jpg")""",
        "Python — Clasificación de Texturas (procesamiento píxel a píxel con OpenCV)"
    )

    titulo_h2("4.4 Reglas de Clasificación — Tabla Resumen")
    tbl3 = doc.add_table(rows=6, cols=4)
    tbl3.style = 'Table Grid'
    enc3 = ["Superficie", "Rango H (Tono)", "Saturación / Valor", "Varianza"]
    reg3 = [
        ("Césped",  "30° – 90°",  "S ≥ 40",           "Var > 5"),
        ("Tierra",  "5° – 30°",   "S ≥ 30",           "Var > 15"),
        ("Cemento", "cualquiera", "S < 30, V ≥ 100",  "Var ≤ 15"),
        ("Asfalto", "cualquiera", "S < 35, V < 100",  "Var > 15"),
        ("Otros",   "—",          "No clasificado",   "—"),
    ]
    row0 = tbl3.rows[0]
    cell_bg(row0.cells[0], '15803D'); cell_bg(row0.cells[1], '15803D')
    cell_bg(row0.cells[2], '15803D'); cell_bg(row0.cells[3], '15803D')
    for ci, t in enumerate(enc3):
        r = row0.cells[ci].paragraphs[0].add_run(t)
        r.font.name='Calibri'; r.font.size=Pt(10); r.bold=True; r.font.color.rgb=RGBColor(255,255,255)
        set_cell_margins(row0.cells[ci])
    for i, cols in enumerate(reg3):
        row = tbl3.rows[i+1]
        bg = 'F0FDF4' if i % 2 == 0 else 'DCFCE7'
        for ci, val in enumerate(cols):
            cell_bg(row.cells[ci], bg)
            rv = row.cells[ci].paragraphs[0].add_run(val)
            rv.font.name='Calibri'; rv.font.size=Pt(10)
            rv.bold = (ci == 0); rv.font.color.rgb = C_NAVY if ci==0 else C_DARK
            set_cell_margins(row.cells[ci])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    titulo_h2("4.5 Resultados Visuales")
    agregar_figura("imagen_prueba.jpg",           "Figura 1 — Imagen de prueba sintética (imagen_prueba.jpg)")
    agregar_figura("resultado_texturas.jpg","Figura 2 — Mapa de clasificación de texturas generado por el algoritmo")
    agregar_figura("mascara_texturas.jpg",  "Figura 3 — Máscara de texturas generada")

    # ════════════════════════════════════════════════════════════
    # 5. ACTIVIDAD B — FILTRO DE SUAVIZADO
    # ════════════════════════════════════════════════════════════
    doc.add_page_break()
    titulo_h1("ACTIVIDAD B — FILTRO DE SUAVIZADO 3×3", num=5)
    draw_hr()

    titulo_h2("5.1 Descripción del Problema")
    parrafo(
        "Se requiere diseñar un filtro de promedio que recorra la imagen usando ventanas de 3×3 píxeles "
        "para reducir ruido y suavizar transiciones de color. El filtro debe implementarse directamente "
        "sobre la matriz de píxeles sin usar funciones de filtrado predefinidas de OpenCV."
    )

    titulo_h2("5.2 Fundamento Teórico — Convolución Discreta")
    parrafo(
        "El filtro de promedio es un caso especial de convolución discreta donde el kernel es una "
        "matriz de 3×3 con todos sus valores iguales a 1/9:"
    )
    agregar_codigo(
        """# Kernel del filtro de promedio 3×3
# Todos los pesos son iguales: 1/9 ≈ 0.111
Kernel = [[1/9, 1/9, 1/9],
          [1/9, 1/9, 1/9],
          [1/9, 1/9, 1/9]]

# Para cada píxel (x, y):
# nuevo_valor = suma(vecino * peso) para los 9 vecinos
# = (P(x-1,y-1) + P(x,y-1) + P(x+1,y-1) +
#    P(x-1,y)   + P(x,y)   + P(x+1,y)   +
#    P(x-1,y+1) + P(x,y+1) + P(x+1,y+1)) / 9""",
        "Matemáticas — Kernel del Filtro Promedio 3×3"
    )

    titulo_h2("5.3 Implementación del Filtro (píxel a píxel)")
    agregar_codigo(
        """import cv2
import numpy as np

img = cv2.imread('paisaje.jpg')
H, W, C = img.shape

# Array de salida (evitar modificar la fuente mientras leemos)
resultado = np.zeros_like(img)

# Recorrer cada píxel de la imagen (excluyendo bordes)
for y in range(1, H - 1):
    for x in range(1, W - 1):
        for c in range(C):               # Canal R, G y B por separado
            suma = 0
            # Ventana deslizante 3×3
            for ky in range(-1, 2):      # ky: -1, 0, +1
                for kx in range(-1, 2):  # kx: -1, 0, +1
                    suma += img[y + ky, x + kx, c]  # Acumular vecinos
            resultado[y, x, c] = suma // 9           # Promedio de 9 vecinos

# Guardar la imagen suavizada
cv2.imwrite('resultado_suavizado.jpg', resultado)
print("✅ Filtro de suavizado aplicado → resultado_suavizado.jpg")""",
        "Python — Filtro Promedio 3×3 implementado manualmente (sin blur() de OpenCV)"
    )

    titulo_h2("5.4 Versión Optimizada con NumPy (para imágenes grandes)")
    parrafo(
        "La versión con bucles for es correcta pero lenta. Para imágenes grandes se puede vectorizar "
        "el mismo cálculo con NumPy usando slicing de matrices, que produce exactamente el mismo resultado:"
    )
    agregar_codigo(
        """import cv2
import numpy as np

img  = cv2.imread('paisaje.jpg').astype(np.float32)
H, W, C = img.shape
out  = np.zeros_like(img)

# Suma de los 9 vecinos usando slicing (equivalente al doble for)
# Cada [y0:y1, x0:x1] representa un desplazamiento del kernel
out[1:-1, 1:-1] = (
    img[0:-2, 0:-2] + img[0:-2, 1:-1] + img[0:-2, 2:  ] +   # Fila superior
    img[1:-1, 0:-2] + img[1:-1, 1:-1] + img[1:-1, 2:  ] +   # Fila central
    img[2:  , 0:-2] + img[2:  , 1:-1] + img[2:  , 2:  ]     # Fila inferior
) / 9.0

cv2.imwrite('resultado_suavizado.jpg', out.astype(np.uint8))
print("✅ Filtro optimizado aplicado → resultado_suavizado.jpg")""",
        "Python — Filtro 3×3 vectorizado con NumPy (mismo resultado, 100× más rápido)"
    )

    titulo_h2("5.5 Resultados Visuales — Antes y Después")
    agregar_figura("imagen_prueba.jpg",           "Figura 4 — Imagen original con ruido (antes del filtro)")
    agregar_figura("resultado_suavizado.jpg","Figura 5 — Imagen después del filtro de suavizado 3×3")

    titulo_h2("5.6 Análisis del Efecto")
    parrafo(
        "El filtro de promedio 3×3 reduce el ruido de alta frecuencia al reemplazar cada píxel con "
        "el promedio de su vecindad. El efecto visible es un suavizado o desenfoque leve, especialmente "
        "notorio en los bordes y en zonas con ruido. Si se necesita más suavizado, el filtro puede "
        "aplicarse múltiples veces o aumentar la ventana a 5×5 o 7×7."
    )

    # ════════════════════════════════════════════════════════════
    # 6. ACTIVIDAD C — COVER LA VACA LOLA (VERSIÓN ROCKERA EN HD)
    # ════════════════════════════════════════════════════════════
    doc.add_page_break()
    titulo_h1("ACTIVIDAD C — COVER DE 'LA VACA LOLA' (ROCK VERSION)", num=6)
    draw_hr()

    titulo_h2("6.1 Descripción")
    parrafo(
        "Se realizó una producción multimedia en alta definición (HD 720p, 1280x720) del cover de la canción infantil 'La Vaca Lola', "
        "con una caracterización de estrella de rock ('Rocker Cow'). El proyecto vincula un audio neural sintetizado con un videoclip animado "
        "en el que los movimientos del personaje (balanceo rítmico, gesticulación de boca y rasgueo de guitarra eléctrica) se sincronizan "
        "con los compases y letra de la música sobre un escenario de concierto."
    )

    titulo_h2("6.2 Pipeline de Producción")
    parrafo("El proceso se dividió en tres etapas:")
    bullet("Etapa 1 — Síntesis de Voz (TTS): Generación del audio de la canción con voz neural boliviana.")
    bullet("Etapa 2 — Creación del Escenario y Animación: Dibujo y animación en OpenCV a 1280x720 píxeles.")
    bullet("Etapa 3 — Composición y Mezcla con FFmpeg: Fusión del video de animación con el audio MP3 en un contenedor MP4 H.264 compatible con la web.")

    titulo_h2("6.3 Etapa 1 — Síntesis de Voz con edge-tts")
    parrafo(
        "Se utilizó la librería edge-tts (Microsoft Edge Neural TTS) configurada con la voz neural en español de Bolivia "
        "('es-BO-MarceloNeural'), aplicando un ajuste de velocidad (+5%) y tono (+3Hz) para encajar con la energía del estilo rockero:"
    )
    agregar_codigo(
        """import edge_tts
import asyncio

LETRA = \"\"\"
Looooooola, la vaca Loooolaaaa...
Tiene cabeza y tiene cola, pero hay tristeza en sus ojos...
\"\"\"

async def generar_voz():
    communicate = edge_tts.Communicate(
        text=LETRA,
        voice="es-BO-MarceloNeural",   # Voz neural boliviana
        rate="+5%",                    # Velocidad
        pitch="+3Hz"                   # Tono
    )
    await communicate.save("vaca_lola.mp3")
    print("✅ Audio generado -> vaca_lola.mp3")

asyncio.run(generar_voz())""",
        "Python — Síntesis de voz con edge-tts (Voz boliviana MarceloNeural)"
    )

    titulo_h2("6.4 Etapa 2 — Creación de la Animación en OpenCV (HD 720p)")
    parrafo(
        "La animación se genera cuadro a cuadro a una resolución de 1280x720 y 30 FPS. Se implementaron los siguientes elementos y movimientos:"
    )
    bullet("Escenario de Rock: Un fondo oscuro con truss metálico de iluminación superior y haces de luces semitransparentes de colores (magenta, cian, amarillo, verde) oscilando con transparencia mediante cv2.addWeighted.")
    bullet("Vaca Rockera: Caracterizada con gafas de sol oscuras de aviador con brillo de cristal, collar de púas blancas, peinado tipo punk/mohicano negro y una guitarra eléctrica roja colgada al pecho.")
    bullet("Sincronización Rítmica: La boca oscila en tamaño según el karaoke, el cuerpo se balancea de izquierda a derecha (bobbing) y la pata delantera derecha realiza un movimiento de arriba a abajo simulando el rasgueo de las cuerdas de la guitarra, acelerando la frecuencia del movimiento en los compases del coro.")
    agregar_codigo(
        """# Lógica de dibujo del escenario y rasgueo de la guitarra
for t in range(total_frames):
    # Luces del escenario oscilando
    ang_osc = 0.5 * np.sin(t * 0.05 + f_offset)
    target_x = int(fx + 300 * np.sin(ang_osc))
    # Dibujar haces en un overlay y mezclar con transparencia
    cv2.addWeighted(overlay, 0.22, frame, 0.78, 0, frame)

    # Animación de rasgueo en el cuerpo de la guitarra
    strum_y = gy + 5 + int(30 * np.sin(t * 0.45))
    cv2.line(frame, (cx - 30, cy - 20), (gx - 15, strum_y), (255, 255, 255), 18)  # Pata rasgueando""",
        "Python/OpenCV — Animación de luces de escenario y rasgueo de guitarra"
    )

    titulo_h2("6.5 Etapa 3 — Composición y Codificación con FFmpeg")
    parrafo(
        "FFmpeg se encargó de fusionar el video raw y el audio MP3 en un archivo de salida final, codificando el video en H.264 (perfil yuv420p) "
        "y el audio en AAC, garantizando compatibilidad absoluta con reproductores multimedia de navegadores web:"
    )
    agregar_codigo(
        """cmd = [
    ffmpeg_path, "-y",
    "-i", "vaca_lola_video.mp4",        # Video de animación OpenCV
    "-i", "vaca_lola.mp3",              # Audio sintetizado
    "-c:v", "libx264",                  # Códec H.264 compatible web
    "-pix_fmt", "yuv420p",              # Formato de píxeles requerido por HTML5
    "-c:a", "aac",                      # Códec de audio AAC
    "-shortest",
    "GutierrezCondori_vaca_lola_final.mp4"
]
subprocess.run(cmd, check=True)""",
        "Python — Llamada a FFmpeg para mezcla H.264/AAC"
    )

    titulo_h2("6.6 Resultado Final")
    parrafo(
        "El resultado final es el archivo GutierrezCondori_vaca_lola_final.mp4 (copiado automáticamente a web/vaca_lola_video.mp4), "
        "el cual presenta la animación de la vaca rockera cantando a resolución HD 720p y con audio perfectamente sincronizado. "
        "Este video se reproduce directamente mediante el reproductor personalizado en la plataforma web."
    )

    # ════════════════════════════════════════════════════════════
    # 7. PLATAFORMA WEB
    # ════════════════════════════════════════════════════════════
    doc.add_page_break()
    titulo_h1("PLATAFORMA WEB INTERACTIVA", num=7)
    draw_hr()

    titulo_h2("7.1 Arquitectura del Sitio")
    parrafo(
        "La plataforma web fue desarrollada con tecnologías nativas del navegador (sin frameworks), "
        "priorizando la compatibilidad y el rendimiento:"
    )
    tbl4 = doc.add_table(rows=4, cols=2)
    tbl4.style = 'Table Grid'
    for i, (archivo, descripcion) in enumerate([
        ("index.html", "Estructura HTML5 con 3 secciones (una por actividad), navegación fija, hero y footer"),
        ("style.css",  "Estilos premium con tema verde oscuro, animaciones CSS, diseño responsivo"),
        ("app.js",     "Lógica JavaScript: algoritmos de procesamiento, Three.js para partículas 3D, carga de imágenes"),
    ]):
        row = tbl4.rows[i]
        bg  = 'F0FDF4' if i % 2 == 0 else 'DCFCE7'
        cell_bg(row.cells[0], bg); cell_bg(row.cells[1], bg)
        r1 = row.cells[0].paragraphs[0].add_run(archivo)
        r1.font.name='Calibri'; r1.font.size=Pt(10); r1.bold=True; r1.font.color.rgb=C_NAVY
        r2 = row.cells[1].paragraphs[0].add_run(descripcion)
        r2.font.name='Calibri'; r2.font.size=Pt(10); r2.font.color.rgb=C_DARK
        set_cell_margins(row.cells[0]); set_cell_margins(row.cells[1])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    titulo_h2("7.2 Procesamiento en Vivo — Canvas API")
    parrafo(
        "La sección más técnica de la web permite ejecutar ambos algoritmos (clasificación y suavizado) "
        "directamente en el navegador, sin enviar datos a ningún servidor. Esto es posible gracias al "
        "Canvas API de HTML5 que expone la matriz de píxeles de cualquier imagen:"
    )
    agregar_codigo(
        """// Obtener acceso a los píxeles de la imagen en el Canvas
const canvas = document.getElementById('canvasIn');
const ctx    = canvas.getContext('2d');
ctx.drawImage(imagenCargada, 0, 0, canvas.width, canvas.height);

// getImageData devuelve un array plano [R,G,B,A, R,G,B,A, ...]
const imageData = ctx.getImageData(0, 0, canvas.width, canvas.height);
const data      = imageData.data;  // Uint8ClampedArray

// Acceder a un píxel en posición (x, y):
const i = (y * canvas.width + x) * 4;
const r = data[i];     // Canal Rojo   (0–255)
const g = data[i + 1]; // Canal Verde  (0–255)
const b = data[i + 2]; // Canal Azul   (0–255)
const a = data[i + 3]; // Alpha/Opacidad (255 = opaco)

// Modificar el píxel y escribir el resultado:
const out = ctx.createImageData(canvas.width, canvas.height);
out.data[i]     = nuevoR;
out.data[i + 1] = nuevoG;
out.data[i + 2] = nuevoB;
out.data[i + 3] = 255;
ctx.putImageData(out, 0, 0); // Dibujar el resultado""",
        "JavaScript — Acceso y modificación de píxeles con Canvas API"
    )

    # ════════════════════════════════════════════════════════════
    # 8. REPOSITORIO GITHUB
    # ════════════════════════════════════════════════════════════
    titulo_h1("REPOSITORIO DE CÓDIGO EN GITHUB", num=8)
    draw_hr()
    parrafo(
        "Todo el código del proyecto fue publicado en un repositorio GitHub con documentación básica de "
        "instalación. El repositorio incluye:"
    )
    bullet("README.md — Instrucciones de instalación y ejecución de cada script.")
    bullet("requirements.txt — Lista de dependencias Python (opencv-python, numpy, edge-tts, Pillow).")
    bullet("procesamiento/ — Scripts Python de clasificación de texturas y filtro de suavizado.")
    bullet("web/ — Archivos HTML, CSS y JavaScript de la plataforma web.")
    bullet("generar_informe.py — Script para generar este informe técnico en formato Word.")
    bullet("Imágenes de resultado: resultado_texturas.jpg, resultado_suavizado.jpg, mascara_texturas.jpg.")

    agregar_codigo(
        """# Comandos para instalar y ejecutar el proyecto desde GitHub

# 1. Clonar el repositorio
git clone https://github.com/usuario/ProyectoMultimedia.git
cd ProyectoMultimedia

# 2. Instalar dependencias Python
pip install -r requirements.txt
# Contenido de requirements.txt:
# opencv-python
# numpy
# Pillow
# edge-tts
# python-docx

# 3. Ejecutar clasificación de texturas
python procesamiento/clasificar_texturas.py

# 4. Ejecutar filtro de suavizado
python procesamiento/filtro_suavizado.py

# 5. Generar el cover de La Vaca Lola
python procesamiento/generar_cover.py

# 6. Abrir la plataforma web en el navegador
# Simplemente abrir web/index.html con cualquier navegador""",
        "Bash — Instalación y ejecución del proyecto desde GitHub"
    )

    # ════════════════════════════════════════════════════════════
    # 9. CONCLUSIONES
    # ════════════════════════════════════════════════════════════
    doc.add_page_break()
    titulo_h1("CONCLUSIONES", num=9)
    draw_hr()

    parrafo(
        "El presente trabajo individual demostró que es posible implementar algoritmos completos de "
        "procesamiento de imágenes operando directamente a nivel de píxel, sin depender de funciones "
        "de alto nivel. A continuación las conclusiones por actividad:"
    )
    doc.add_paragraph()

    for i, (titulo, texto) in enumerate([
        (
            "Clasificación de Texturas",
            "La combinación del espacio HSV con la varianza local de brillo (5×5) permite diferenciar "
            "satisfactoriamente superficies de césped, tierra, cemento y asfalto en condiciones de "
            "iluminación variada. El enfoque es análogo a los clasificadores de imágenes satelitales."
        ),
        (
            "Filtro de Suavizado 3×3",
            "El filtro de promedio con ventana deslizante de 3×3 reduce el ruido de alta frecuencia "
            "de forma efectiva. La implementación manual refuerza la comprensión de la convolución discreta, "
            "concepto fundamental en el procesamiento de imágenes y en las redes neuronales convolucionales."
        ),
        (
            "Cover de La Vaca Lola",
            "La integración de edge-tts para síntesis de voz neural con FFmpeg para la composición "
            "de video permitió producir un cover multimedia de calidad profesional, con sincronización "
            "correcta entre el audio y la animación del personaje."
        ),
        (
            "Plataforma Web",
            "La implementación de los algoritmos directamente en JavaScript con la Canvas API permite "
            "que el usuario ejecute el procesamiento en tiempo real sin necesidad de instalaciones, "
            "haciendo el proyecto más accesible y demostrable."
        ),
    ], start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent  = Cm(0.5)
        r1 = p.add_run(f"{i}. {titulo}: ")
        r1.font.name='Calibri'; r1.font.size=Pt(10.5); r1.bold=True; r1.font.color.rgb=C_NAVY
        r2 = p.add_run(texto)
        r2.font.name='Calibri'; r2.font.size=Pt(10.5); r2.font.color.rgb=C_DARK

    # ════════════════════════════════════════════════════════════
    # 10. REFERENCIAS
    # ════════════════════════════════════════════════════════════
    titulo_h1("REFERENCIAS", num=10)
    draw_hr()
    for ref in [
        "Gonzalez, R. C., & Woods, R. E. (2018). Digital Image Processing (4th ed.). Pearson.",
        "OpenCV Documentation. (2024). Image Processing in OpenCV. https://docs.opencv.org/",
        "Microsoft. (2024). Edge TTS — Neural Text-to-Speech. https://github.com/rany2/edge-tts",
        "FFmpeg. (2024). FFmpeg Documentation. https://ffmpeg.org/documentation.html",
        "MDN Web Docs. (2024). CanvasRenderingContext2D.getImageData(). https://developer.mozilla.org/",
        "Three.js. (2024). Three.js Documentation r128. https://threejs.org/docs/",
    ]:
        bullet(ref)

    # ════════════════════════════════════════════════════════════
    # GUARDAR
    # ════════════════════════════════════════════════════════════
    output = "Informe_Tecnico_JonathanGutierrez.docx"
    doc.save(output)
    print(f"\n✅ Informe generado exitosamente: {output}")
    return output


if __name__ == "__main__":
    crear_informe()
